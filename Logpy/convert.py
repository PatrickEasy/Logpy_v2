"""
Convert between file formats without pandoc.

Supported conversions:
    .md   -> .docx
    .csv  -> .json
    .json -> .csv

Test-file generators (write to the ``tfiles/`` folder):
    .md   (markdown exercising every element ``md_to_docx`` handles)
    .docx (direct python-docx output, useful for round-trip testing)
    .csv  (random-word rows)
    .json (list of flat objects with the occasional nested value)

CLI usage:
    python convert.py test1.md               # writes test1.docx
    python convert.py data.csv               # writes data.json
    python convert.py data.json              # writes data.csv
    python convert.py data.csv out.json      # explicit output path

    python convert.py --gen sample.md        # writes tfiles/sample.md
    python convert.py --gen sample.docx
    python convert.py --gen sample.csv
    python convert.py --gen sample.json

    python convert.py --batch ./docs .md     # convert every .md under ./docs

Install once:
    pip install markdown python-docx beautifulsoup4

Every conversion and generator call is wrapped in a
:class:`Logpy.Timer` and reported through :func:`Logpy.printtime`, so
running this script also produces a structured JSON log in ``logs/``.
"""

import csv
import json
import random
import sys
from pathlib import Path

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

# Shared Logpy utilities — use relative imports when loaded as a package,
# fall back to absolute when running this file directly.
try:
    from .print_utils import smart_print, ok, info, err, C
    from .printtime import printtime, find_files_with_extension
    from .timer import Timer
except ImportError:
    from print_utils import smart_print, ok, info, err, C
    from printtime import printtime, find_files_with_extension
    from timer import Timer # type: ignore


# ---------------------------------------------------------------------------
# Markdown -> DOCX
# ---------------------------------------------------------------------------

def _add_runs(paragraph, node):
    """Walk an HTML node and append runs to a docx paragraph, preserving bold/italic/code."""
    for child in node.children:
        if getattr(child, "name", None) is None:
            paragraph.add_run(str(child))
        elif child.name in ("strong", "b"):
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name in ("em", "i"):
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == "code":
            run = paragraph.add_run(child.get_text())
            run.font.name = "Consolas"
        elif child.name == "a":
            run = paragraph.add_run(child.get_text())
            run.underline = True
        else:
            _add_runs(paragraph, child)


def md_to_docx(src: Path, dst: Path) -> None:
    """Convert a markdown file to a Word document."""
    html = markdown.markdown(
        src.read_text(encoding="utf-8"),
        extensions=["fenced_code", "tables"],
    )
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    for el in soup.children:
        name = getattr(el, "name", None)
        if name is None:
            continue
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            doc.add_heading(el.get_text(), level=int(name[1]))
        elif name == "p":
            p = doc.add_paragraph()
            _add_runs(p, el)
        elif name in ("ul", "ol"):
            style = "List Bullet" if name == "ul" else "List Number"
            for li in el.find_all("li", recursive=False):
                p = doc.add_paragraph(style=style)
                _add_runs(p, li)
        elif name == "pre":
            code_text = el.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif name == "blockquote":
            p = doc.add_paragraph(style="Intense Quote")
            _add_runs(p, el)
        elif name == "hr":
            doc.add_paragraph("_" * 40)
        elif name == "table":
            rows = el.find_all("tr")
            if not rows:
                continue
            cols = max(len(r.find_all(["td", "th"])) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Light Grid Accent 1"
            for r_idx, row in enumerate(rows):
                for c_idx, cell in enumerate(row.find_all(["td", "th"])):
                    table.cell(r_idx, c_idx).text = cell.get_text()

    doc.save(dst)


# ---------------------------------------------------------------------------
# CSV <-> JSON
# ---------------------------------------------------------------------------

def csv_to_json(src: Path, dst: Path) -> None:
    """Convert a CSV file to a JSON list-of-objects."""
    with src.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
    with dst.open("w", encoding="utf-8") as f:
        json.dump(rows, f, indent=2, ensure_ascii=False)


def json_to_csv(src: Path, dst: Path) -> None:
    """Convert a JSON file (list-of-objects or single object) to CSV."""
    data = json.loads(src.read_text(encoding="utf-8"))

    # Accept either a list of dicts, or a single dict (wrap it).
    if isinstance(data, dict):
        data = [data]
    if not isinstance(data, list) or not all(isinstance(r, dict) for r in data):
        raise ValueError(
            "JSON must be a list of flat objects (or a single object). "
            "Nested structures aren't supported — flatten first."
        )

    # Union of keys across all rows, preserving first-seen order.
    fieldnames: list[str] = []
    seen = set()
    for row in data:
        for key in row.keys():
            if key not in seen:
                seen.add(key)
                fieldnames.append(key)

    with dst.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in data:
            # Stringify nested values so csv doesn't choke.
            writer.writerow({
                k: (json.dumps(v, ensure_ascii=False) if isinstance(v, (dict, list)) else v)
                for k, v in row.items()
            })


# ---------------------------------------------------------------------------
# Test-file generators
# ---------------------------------------------------------------------------

# A small but varied word pool. Kept inline so the script has no extra deps.
_WORDS = (
    "alpha bravo charlie delta echo foxtrot golf hotel india juliet kilo lima "
    "mike november oscar papa quebec romeo sierra tango uniform victor whiskey "
    "xray yankee zulu apple banana cherry dragon elephant falcon glacier harbor "
    "island jungle kernel lantern meadow nebula ocean prairie quasar river "
    "summit tundra umbra valley willow xenon yonder zephyr cobalt crimson "
    "azure scarlet saffron jade ivory onyx pearl amber violet teal "
    "orbit photon quantum vector matrix tensor entropy plasma fusion lattice"
).split()


def _word() -> str:
    return random.choice(_WORDS)


def _words(n: int) -> str:
    return " ".join(_word() for _ in range(n))


def _sentence(min_words: int = 6, max_words: int = 14) -> str:
    s = _words(random.randint(min_words, max_words))
    return s.capitalize() + "."


def _paragraph(sentences: int = 4) -> str:
    return " ".join(_sentence() for _ in range(sentences))


def generate_md(dst: Path, sections: int = 4) -> None:
    """Write a markdown file that exercises every element md_to_docx handles."""
    lines: list[str] = [f"# {_words(3).title()}", "", _paragraph(), ""]
    for i in range(sections):
        lines += [f"## Section {i + 1}: {_words(2).title()}", "", _paragraph(3), ""]

        # Bullet list with inline formatting.
        lines.append("Some notable points:")
        lines.append("")
        for _ in range(random.randint(3, 5)):
            lines.append(
                f"- **{_word()}** and *{_word()}* with `{_word()}_code`"
            )
        lines.append("")

        # Numbered list.
        lines.append("Steps to reproduce:")
        lines.append("")
        for step in range(1, random.randint(3, 5) + 1):
            lines.append(f"{step}. {_sentence()}")
        lines.append("")

        # Blockquote + code block on alternating sections.
        if i % 2 == 0:
            lines += [f"> {_sentence()}", ""]
        else:
            lines += ["```python", f"def {_word()}():", f"    return '{_word()}'", "```", ""]

    # A small table at the end.
    lines += [
        "| Name | Value | Notes |",
        "| --- | --- | --- |",
        *[f"| {_word()} | {random.randint(1, 999)} | {_words(3)} |" for _ in range(4)],
        "",
    ]

    dst.write_text("\n".join(lines), encoding="utf-8")


def generate_docx(dst: Path, sections: int = 4) -> None:
    """Write a .docx directly (not via md_to_docx) so round-trips can be tested."""
    doc = Document()
    doc.add_heading(_words(3).title(), level=1)
    doc.add_paragraph(_paragraph())

    for i in range(sections):
        doc.add_heading(f"Section {i + 1}: {_words(2).title()}", level=2)
        doc.add_paragraph(_paragraph(3))

        for _ in range(random.randint(3, 5)):
            doc.add_paragraph(_sentence(), style="List Bullet")
        for _ in range(random.randint(3, 5)):
            doc.add_paragraph(_sentence(), style="List Number")

        if i % 2 == 0:
            doc.add_paragraph(_sentence(), style="Intense Quote")

    doc.save(dst)


def generate_csv(dst: Path, rows: int = 20, cols: int = 5) -> None:
    """Write a CSV with random header and mixed word/number cells."""
    # Guarantee unique column names.
    header = random.sample(_WORDS, k=cols)
    with dst.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(header)
        for _ in range(rows):
            writer.writerow([
                _word() if random.random() < 0.6 else random.randint(0, 9999)
                for _ in range(cols)
            ])


def generate_json(dst: Path, rows: int = 20, cols: int = 5) -> None:
    """Write a JSON list-of-objects. A few rows include nested values on purpose
    so json_to_csv's fallback-to-string path gets exercised."""
    keys = random.sample(_WORDS, k=cols)
    data = []
    for i in range(rows):
        row = {
            k: (_word() if random.random() < 0.6 else random.randint(0, 9999))
            for k in keys
        }
        # Sprinkle in nested values on ~1 in 5 rows.
        if i % 5 == 4:
            row[keys[0]] = {"tag": _word(), "score": random.randint(0, 100)}
        data.append(row)
    dst.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


GENERATORS = {
    ".md": generate_md,
    ".docx": generate_docx,
    ".csv": generate_csv,
    ".json": generate_json,
}

# All generated test files land in this folder (created if missing).
TFILES_DIR = Path("tfiles")


def generate(dst_path: str) -> Path:
    """Generate a test file whose type is inferred from the extension.

    The file is always written to :data:`TFILES_DIR` (``tfiles/``), which
    is created if missing. Only the filename portion of ``dst_path`` is
    used, so bare names and full paths both resolve there.
    """
    # Take only the filename so bare names and full paths both end up in tfiles/.
    dst = TFILES_DIR / Path(dst_path).name
    ext = dst.suffix.lower()
    if ext not in GENERATORS:
        raise ValueError(f"No generator for {ext!r}. Supported: {list(GENERATORS)}")
    TFILES_DIR.mkdir(parents=True, exist_ok=True)
    with Timer(f"convert.generate[{ext}]"):
        GENERATORS[ext](dst)
    printtime(f"Generated test file: {dst}")
    return dst


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

# Map (input_ext, output_ext) -> converter function.
CONVERTERS = {
    (".md", ".docx"): md_to_docx,
    (".csv", ".json"): csv_to_json,
    (".json", ".csv"): json_to_csv,
}

# If only input is given, pick the default output extension.
DEFAULT_OUTPUT = {
    ".md": ".docx",
    ".csv": ".json",
    ".json": ".csv",
}


def convert(src_path: str, dst_path: str | None = None) -> Path:
    """Convert ``src_path`` to ``dst_path``.

    If ``dst_path`` is omitted the output extension is chosen from
    :data:`DEFAULT_OUTPUT`. Raises :class:`FileNotFoundError` if the
    source does not exist, or :class:`ValueError` for unsupported
    conversions.
    """
    src = Path(src_path)
    if not src.exists():
        raise FileNotFoundError(src)

    src_ext = src.suffix.lower()
    if dst_path:
        dst = Path(dst_path)
    else:
        if src_ext not in DEFAULT_OUTPUT:
            raise ValueError(f"No default output format for {src_ext!r}")
        dst = src.with_suffix(DEFAULT_OUTPUT[src_ext])

    dst_ext = dst.suffix.lower()
    key = (src_ext, dst_ext)
    if key not in CONVERTERS:
        supported = ", ".join(f"{a}->{b}" for a, b in CONVERTERS)
        raise ValueError(f"Unsupported conversion {src_ext}->{dst_ext}. Supported: {supported}")

    with Timer(f"convert[{src_ext}->{dst_ext}]"):
        CONVERTERS[key](src, dst)
    printtime(f"Converted {src} -> {dst}")
    return dst


def batch_convert(directory: str, extension: str, recursive: bool = True) -> list[Path]:
    """Convert every file under ``directory`` with the given ``extension``.

    Uses :func:`Logpy.find_files_with_extension` to discover inputs, then
    calls :func:`convert` on each. Returns the list of output paths.
    """
    inputs = find_files_with_extension(directory, extension, recursive=recursive)
    if not inputs:
        smart_print(f"No {extension} files found in {directory}", "info")
        return []

    smart_print(f"Batch converting {len(inputs)} {extension} file(s) in {directory}", "title")
    outputs: list[Path] = []
    with Timer("convert.batch", auto_log=False) as batch:
        for src in inputs:
            try:
                outputs.append(convert(src))
                ok(f"{src}")
            except Exception as exc:
                smart_print(f"Failed to convert {src}: {exc}", "info")
    printtime(
        f"Batch converted {len(outputs)}/{len(inputs)} file(s) in "
        f"{Timer._format_time(batch.elapsed())}"
    )
    return outputs


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _print_usage() -> None:
    smart_print("Usage", "title")
    print("  python convert.py <input> [output]        # convert")
    print("  python convert.py --gen <path>            # generate test file")
    print("  python convert.py --batch <dir> <ext>     # convert every file with <ext>")
    print("Supported: .md -> .docx, .csv <-> .json")


def cli(argv: list[str] | None = None) -> None:
    """Command-line entry point.

    Exposed as ``logpy-convert`` when the package is pip-installed, and
    invoked by ``python -m Logpy.convert`` / ``python convert.py``.
    """
    args = list(sys.argv[1:] if argv is None else argv)
    if not args:
        _print_usage()
        sys.exit(1)

    try:
        if args[0] in ("--gen", "--generate"):
            if len(args) < 2:
                smart_print("Usage: python convert.py --gen <path>", "info")
                sys.exit(1)
            out = generate(args[1])
            ok(f"Generated {out}")
        elif args[0] in ("--batch",):
            if len(args) < 3:
                smart_print("Usage: python convert.py --batch <dir> <ext>", "info")
                sys.exit(1)
            directory, extension = args[1], args[2]
            outputs = batch_convert(directory, extension)
            ok(f"Batch complete: {len(outputs)} file(s) written")
        else:
            filename = args[0]
            output = args[1] if len(args) > 1 else None
            out = convert(filename, output)
            ok(f"Wrote {out}")
    except (FileNotFoundError, ValueError) as exc:
        # err() prints the message and exits 1 — matches the old behaviour.
        err(str(exc))


if __name__ == "__main__":
    cli()
